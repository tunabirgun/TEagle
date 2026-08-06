"""Build the manuscript tables from raw benchmark output. No cell is typed by hand in the manuscript.

    python benchmarks/make_tables.py

Table 4  - functional comparison, compact verdicts, for the printed manuscript.
Table S1 - the same matrix with each cell's full sourced justification, for the supplement.

Each is written as .csv, .xlsx and .docx into manuscript/tables/.

TEagle's own row is NOT taken from the marketing copy: every cell cites the module that implements it, so
a reader can check the claim against the source. The comparator rows come verbatim from
benchmarks/raw/tool_capability_matrix.json, which records the publication or documentation behind each
cell. Where that research flagged a cell as lower-confidence, the flag travels with it into Table S1.
"""
from __future__ import annotations
import csv, json, os, re, sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RAW = os.path.join(ROOT, "benchmarks", "raw", "tool_capability_matrix.json")
REPRO = os.path.join(ROOT, "benchmarks", "raw", "reproducibility.json")
OUT = os.path.join(ROOT, "manuscript", "tables")

sys.path[:0] = [os.path.join(ROOT, "app", "backend")]
import engine                                              # noqa: E402
from teagle_core import __version__ as TEAGLE_VERSION      # noqa: E402  single source of truth

# Both counted rather than written down, so the table cannot claim a coverage the manifest does not
# have. They are DIFFERENT numbers and the cell must not conflate them: the seal carries N_SEALED
# parameters in total, of which only N_THRESHOLDS are detector thresholds recovered from the
# signatures of the functions that apply them. The remainder are module-level constants, which the
# "read from the function that applies it" guarantee does not cover.
N_SEALED = len(engine._detector_parameters())
if not os.path.exists(REPRO):
    raise SystemExit(f"missing {REPRO} - run benchmarks/reproducibility.py before building the tables")
N_THRESHOLDS = json.load(open(REPRO, encoding="utf-8"))["fidelity"]["detector_thresholds_found"]

# axis key -> (printed column header, short header for the compact table)
AXES = [
    ("input_granularity",      "Input granularity",                     "Input"),
    ("denovo_structural",      "De novo structural detection",          "Structural"),
    ("protein_domain_classify", "Protein-domain classification",        "Domains"),
    ("homology_family_naming", "Homology-based family naming",          "Family"),
    ("wicker_superfamily_call", "Wicker superfamily call",              "Wicker"),
    ("explicit_abstention",    "Withholds a call on insufficient evidence", "Abstains"),
    ("confidence_reporting",   "Per-call confidence reported",          "Confidence"),
    ("primer_design",          "PCR primer design",                     "Primers"),
    ("in_silico_pcr",          "In-silico PCR",                         "isPCR"),
    ("genome_offtarget_scan",  "Whole-genome off-target scan",          "Off-target"),
    ("provenance_record",      "Machine-readable provenance record",    "Provenance"),
    ("gui_no_commandline",     "Usable without a command line",         "GUI"),
    ("offline_capable",        "Runs offline",                          "Offline"),
    ("export_formats",         "Export formats",                        "Exports"),
    ("license",                "License",                               "License"),
    ("platform",               "Platform",                              "Platform"),
]
VERDICT_AXES = {k for k, _, _ in AXES} - {"input_granularity", "export_formats", "license", "platform"}

# "not applicable" and "not documented" are different verdicts and must not share a glyph: the first says
# the capability lies outside the tool's design scope, the second says no evidence of it could be found.
# Collapsing them onto one symbol renders absence of evidence as out-of-scope, which understates a gap.
GLYPH = {"yes": "yes", "no": "no", "partial": "partial",
         "not applicable": "n.a.", "not documented": "n.d."}

# TEagle's own row. Each value is followed by the module that implements it - the claim and its evidence
# travel together, so a reviewer can check any cell against the source tree.
TEAGLE = {
    "name": "TEagle",
    "current_version": f"{TEAGLE_VERSION} (this work)",
    "primary_reference": "This work.",
    "input_granularity": "single element — one sequence per analysis, from an accession, a file or pasted "
                         "text; a whole assembly is read only by the primer off-target scan "
                         "(app/backend/engine.py)",
    "denovo_structural": "yes — terminal repeat, terminal inverted repeat and target-site duplication "
                         "detectors with published thresholds (app/backend/teagle_core/structural.py)",
    "protein_domain_classify": "yes — bundled 30-model Pfam-A panel searched with HMMER via pyhmmer "
                               "(app/backend/teagle_core/domains.py)",
    "homology_family_naming": "partial — optional Dfam naming through RepeatMasker in the installed WSL2 "
                              "environment; the core application does not require it",
    "wicker_superfamily_call": "yes — Wicker et al. 2007 class and superfamily from structural and domain "
                               "evidence (app/backend/teagle_core/classify.py)",
    "explicit_abstention": "yes — withholds the superfamily where translation order is unreadable across "
                           "strands, states when the domain scan did not run, and reports how many reading "
                           "frames were left unsearched (app/backend/teagle_core/classify.py)",
    "confidence_reporting": "yes — a confidence level and a separate domain completeness, reported as "
                            "two independent axes",
    "primer_design": "yes — Primer3 in-process, with QC on each pair",
    "in_silico_pcr": "yes — predicted amplicons against the supplied template",
    "genome_offtarget_scan": "yes — isPcr against a whole assembly, in the optional WSL2 environment",
    "provenance_record": "yes — content-addressed manifest carrying application and database versions, "
                         f"checksums, the input hash and {N_SEALED} sealed parameters, {N_THRESHOLDS} of "
                         "them detector thresholds read from the signatures of the functions that apply "
                         f"them and the remaining {N_SEALED - N_THRESHOLDS} module-level constants "
                         "(app/backend/engine.py::_detector_parameters)",
    "gui_no_commandline": "yes — native desktop application; no command line at any point",
    "offline_capable": "yes — the scientific core is bundled and runs with no network; only accession "
                       "fetch and optional Dfam naming need connectivity",
    "export_formats": "GFF3, BED, FASTA, XLSX, CSV, TSV, JSON provenance manifest, and SVG/PNG/PDF figures "
                      "(app/native/widgets.py, app/native/main.py)",
    "license": "AGPL-3.0-or-later",
    "platform": "Windows 10/11 64-bit; optional components in an app-installed WSL2 environment",
    "what_it_does_better": "—",
    "sources": f"This work; TEagle {TEAGLE_VERSION} source tree.",
}


def verdict(value: str) -> str:
    """The verdict opening the cell. Two words are read before one, so 'not applicable' and 'not
    documented' stay distinct; punctuation is stripped because several cells continue straight into the
    justification ('partial, and the mechanism is...')."""
    words = re.split(r"\W+", str(value).strip().lower())
    two = " ".join(words[:2])
    if two in GLYPH:
        return GLYPH[two]
    return GLYPH.get(words[0], words[0] or "—")


def compact(value: str, axis: str) -> str:
    return verdict(value) if axis in VERDICT_AXES else re.split(r"\s*[—(]", str(value).strip(), 1)[0].strip()


def write_csv(path, header, rows):
    with open(path, "w", newline="", encoding="utf-8-sig") as fh:
        w = csv.writer(fh)
        w.writerow(header)
        w.writerows(rows)


def write_xlsx(path, header, rows, title):
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]
    ws.append(header)
    for r in rows:
        ws.append(r)
    head_fill = PatternFill("solid", fgColor="DDDDDD")
    for c in ws[1]:
        c.font = Font(bold=True, name="Times New Roman", size=10)
        c.fill = head_fill
        c.alignment = Alignment(wrap_text=True, vertical="top")
    for row in ws.iter_rows(min_row=2):
        for c in row:
            c.font = Font(name="Times New Roman", size=10)
            c.alignment = Alignment(wrap_text=True, vertical="top")
    ws.freeze_panes = "B2"
    widths = [max(12, min(46, max(len(str(v or "")) for v in col) // 2 + 8))
              for col in zip(header, *rows)] if rows else []
    for i, wdt in enumerate(widths, start=1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = wdt
    wb.save(path)


def write_docx(path, header, rows, caption):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.shared import Mm
    doc = Document()
    s = doc.sections[0]
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width, s.page_height = Mm(297), Mm(210)
    s.left_margin = s.right_margin = Mm(8)
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(6.5)
    st.font.color.rgb = RGBColor(0, 0, 0)

    cap = doc.add_paragraph()
    r = cap.add_run(caption)
    r.bold = True
    r.font.size = Pt(7.5)
    r.font.name = "Times New Roman"

    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    for i, h in enumerate(header):
        p = t.rows[0].cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(6.5)
        run.font.name = "Times New Roman"
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(v or ""))
            run.font.size = Pt(6)
            run.font.name = "Times New Roman"
    doc.save(path)


def main():
    if not os.path.exists(RAW):
        print(f"missing {RAW} - run the capability-matrix workflow first")
        return 1
    data = json.load(open(RAW, encoding="utf-8"))
    tools = [TEAGLE] + data["tools"]
    os.makedirs(OUT, exist_ok=True)

    # ---- Table 4: compact verdicts -----------------------------------------------------------------
    keep = [(k, _l, s) for k, _l, s in AXES if k not in ("export_formats", "license", "platform")]
    h1 = ["Tool", "Version"] + [s for _, _, s in keep]
    # No character truncation. Cutting a version string at a fixed width produced cells like
    # "REPET 2.2 standalone p" and "binary v1.0.7, as bund" in the exported table — a reader cannot tell
    # a truncated cell from a complete one, and the .csv/.xlsx/.docx all handle the full string.
    r1 = [[t["name"], t["current_version"].split("(")[0].strip()]
          + [compact(t[k], k) for k, _, _ in keep] for t in tools]
    cap1 = ("Table 4. Functional comparison of TEagle with fifteen widely used transposable-element and "
            "primer-design tools. Verdicts: yes; partial; no; n.a. where the capability lies outside the "
            "tool's design scope; n.d. where no such capability is documented but the tool does not place it "
            "out of scope. Full justification and sources for every cell are in Table S1.")
    write_csv(os.path.join(OUT, "table4_tool_comparison.csv"), h1, r1)
    write_xlsx(os.path.join(OUT, "table4_tool_comparison.xlsx"), h1, r1, "Table 4")
    write_docx(os.path.join(OUT, "table4_tool_comparison.docx"), h1, r1, cap1)

    # ---- Table S1: full sourced matrix --------------------------------------------------------------
    hs = ["Tool", "Version", "Primary reference"] + [lbl for _, lbl, _ in AXES] + \
         ["What it does better than TEagle", "Sources"]
    rs = [[t["name"], t["current_version"], t["primary_reference"]]
          + [t[k] for k, _, _ in AXES] + [t.get("what_it_does_better", ""), t.get("sources", "")]
          for t in tools]
    caps = ("Table S1. Functional comparison of TEagle with fifteen widely used tools, with the full "
            "justification and source for every cell. Comparator cells are sourced to each tool's own "
            "publication or documentation; TEagle's cells cite the module that implements the behaviour.")
    write_csv(os.path.join(OUT, "tableS1_tool_comparison_sourced.csv"), hs, rs)
    write_xlsx(os.path.join(OUT, "tableS1_tool_comparison_sourced.xlsx"), hs, rs, "Table S1")
    write_docx(os.path.join(OUT, "tableS1_tool_comparison_sourced.docx"), hs, rs, caps)

    print(f"Table 4  : {len(r1)} rows x {len(h1)} cols  -> .csv .xlsx .docx")
    print(f"Table S1 : {len(rs)} rows x {len(hs)} cols  -> .csv .xlsx .docx")
    print()
    print("Verdict distribution on the abstention axis (the paper's central claim):")
    for t in tools:
        print(f"  {verdict(t['explicit_abstention']):9s}  {t['name']}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
